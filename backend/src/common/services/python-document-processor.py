#!/usr/bin/env python3
"""
Python service for document processing including chunking and Word document creation
Enhanced with advanced NLP processing capabilities
"""

import sys
import json
import argparse
import re
import os
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# Import advanced processor
try:
    from advanced_document_processor import AdvancedDocumentProcessor
    ADVANCED_PROCESSING_AVAILABLE = True
except ImportError:
    ADVANCED_PROCESSING_AVAILABLE = False

class DocumentChunker:
    """Handles document chunking using Python"""
    
    def __init__(self, max_chunk_size: int = 1000, overlap_size: int = 200):
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove page numbers and headers/footers
        text = re.sub(r'^\d+\s*$', '', text, flags=re.MULTILINE)
        # Remove common headers/footers
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'© \d+.*', '', text)
        return text.strip()
    
    def create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create chunks from text using sliding window approach"""
        cleaned_text = self.preprocess_text(text)
        
        # Split into sentences for better chunking
        sentences = re.split(r'(?<=[.!?])\s+', cleaned_text)
        
        chunks = []
        current_chunk = ""
        chunk_index = 0
        start_pos = 0
        
        for sentence in sentences:
            if not sentence.strip():
                continue
                
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > self.max_chunk_size and current_chunk:
                # Save current chunk
                chunk = {
                    "chunkText": current_chunk.strip(),
                    "startPos": start_pos,
                    "endPos": start_pos + len(current_chunk),
                    "tokenCount": len(current_chunk.split()),
                    "metadata": {
                        "chunkIndex": chunk_index,
                        "isHeader": False,
                        "isFooter": False,
                        "isTable": False,
                        "isList": False,
                        "confidence": 1.0,
                        "language": "en"
                    }
                }
                chunks.append(chunk)
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.overlap_size:] if len(current_chunk) > self.overlap_size else current_chunk
                current_chunk = overlap_text + " " + sentence
                start_pos = start_pos + len(current_chunk) - len(overlap_text) - len(sentence)
                chunk_index += 1
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add final chunk
        if current_chunk.strip():
            chunk = {
                "chunkText": current_chunk.strip(),
                "startPos": start_pos,
                "endPos": start_pos + len(current_chunk),
                "tokenCount": len(current_chunk.split()),
                "metadata": {
                    "chunkIndex": chunk_index,
                    "isHeader": False,
                    "isFooter": False,
                    "isTable": False,
                    "isList": False,
                    "confidence": 1.0,
                    "language": "en"
                }
            }
            chunks.append(chunk)
        
        return chunks

class WordDocumentCreator:
    """Creates Word documents from chunks"""
    
    def __init__(self):
        pass
    
    def create_word_document(self, chunks: List[Dict[str, Any]], document_title: str = "Legal Document") -> bytes:
        """Create a Word document from chunks"""
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = document_title
        doc.core_properties.author = "Legal Document System"
        doc.core_properties.subject = "Converted from Legal Document System"
        
        # Add title
        title = doc.add_heading(document_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph("=" * 50)
        
        # Process each chunk
        for i, chunk in enumerate(chunks):
            # Add chunk header
            chunk_header = doc.add_heading(f'Section {i + 1}', level=2)
            
            # Add chunk content
            paragraph = doc.add_paragraph()
            
            # Split text into sentences for better formatting
            sentences = chunk.get('chunkText', '').split('. ')
            for j, sentence in enumerate(sentences):
                if sentence.strip():
                    if j > 0:
                        paragraph.add_run(sentence.strip() + '. ')
                    else:
                        paragraph.add_run(sentence.strip() + '. ')
            
            # Add chunk metadata
            if chunk.get('metadata'):
                metadata = chunk.get('metadata', {})
                if metadata.get('chunkIndex') is not None:
                    meta_para = doc.add_paragraph()
                    meta_para.add_run(f"Chunk Index: {metadata.get('chunkIndex')}").italic = True
                    
                if metadata.get('tokenCount'):
                    meta_para = doc.add_paragraph()
                    meta_para.add_run(f"Token Count: {metadata.get('tokenCount')}").italic = True
            
            # Add spacing between chunks
            doc.add_paragraph()
        
        # Add footer
        doc.add_paragraph("=" * 50)
        footer_para = doc.add_paragraph()
        footer_para.add_run("Generated by Legal Document System").italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to bytes
        import io
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

def process_document(input_data: Dict[str, Any]) -> Dict[str, Any]:
    """Process document: chunk text and create Word document"""
    try:
        # Extract text from input
        text = input_data.get('text', '')
        document_title = input_data.get('title', 'Legal Document')
        file_path = input_data.get('file_path')
        document_id = input_data.get('document_id', 'unknown')
        
        if not text.strip():
            raise ValueError("No text content provided")
        
        # Try advanced processing first if available and file path is provided
        if ADVANCED_PROCESSING_AVAILABLE and file_path and os.path.exists(file_path):
            try:
                advanced_processor = AdvancedDocumentProcessor()
                result = advanced_processor.process_document(file_path, document_id)
                
                if result.success:
                    # Convert advanced chunks to legacy format
                    chunks = []
                    for chunk in result.chunks:
                        legacy_chunk = {
                            "chunkText": chunk['content'],
                            "startPos": chunk['metadata']['start_char'],
                            "endPos": chunk['metadata']['end_char'],
                            "tokenCount": chunk['metadata']['word_count'],
                            "metadata": {
                                "chunkIndex": chunk['metadata']['chunk_index'],
                                "isHeader": False,
                                "isFooter": False,
                                "isTable": False,
                                "isList": False,
                                "confidence": chunk['metadata'].get('confidence_score', 1.0),
                                "language": "en",
                                "keywords": chunk['metadata'].get('keywords', []),
                                "entities": chunk['metadata'].get('entities', []),
                                "semantic_embedding": chunk['metadata'].get('semantic_embedding', [])
                            }
                        }
                        chunks.append(legacy_chunk)
                    
                    # Create Word document
                    word_creator = WordDocumentCreator()
                    word_bytes = word_creator.create_word_document(chunks, document_title)
                    
                    return {
                        "chunks": chunks,
                        "word_document": word_bytes.hex(),
                        "chunk_count": len(chunks),
                        "total_tokens": sum(chunk.get('tokenCount', 0) for chunk in chunks),
                        "processing_method": "advanced",
                        "advanced_metadata": result.metadata
                    }
            except Exception as e:
                print(f"Advanced processing failed, falling back to basic: {e}", file=sys.stderr)
        
        # Fallback to basic processing
        chunker = DocumentChunker(
            max_chunk_size=input_data.get('max_chunk_size', 1000),
            overlap_size=input_data.get('overlap_size', 200)
        )
        
        # Create chunks
        chunks = chunker.create_chunks(text)
        
        # Create Word document
        word_creator = WordDocumentCreator()
        word_bytes = word_creator.create_word_document(chunks, document_title)
        
        # Prepare response
        response = {
            "chunks": chunks,
            "word_document": word_bytes.hex(),  # Convert to hex string for JSON
            "chunk_count": len(chunks),
            "total_tokens": sum(chunk.get('tokenCount', 0) for chunk in chunks),
            "processing_method": "basic"
        }
        
        return response
        
    except Exception as e:
        return {
            "error": str(e),
            "chunks": [],
            "word_document": "",
            "chunk_count": 0,
            "total_tokens": 0
        }

def main():
    """Main function to handle command line arguments"""
    parser = argparse.ArgumentParser(description='Process document: chunk text and create Word document')
    parser.add_argument('--input', required=True, help='JSON string containing document data')
    parser.add_argument('--output', help='Output file path (optional)')
    
    args = parser.parse_args()
    
    try:
        # Parse input data
        input_data = json.loads(args.input)
        
        # Process document
        result = process_document(input_data)
        
        if args.output:
            # Save to file
            with open(args.output, 'w') as f:
                json.dump(result, f, indent=2)
            print(f"Results saved to: {args.output}")
        else:
            # Output to stdout as JSON
            print(json.dumps(result))
            
    except Exception as e:
        error_result = {
            "error": str(e),
            "chunks": [],
            "word_document": "",
            "chunk_count": 0,
            "total_tokens": 0
        }
        print(json.dumps(error_result))
        sys.exit(1)

if __name__ == "__main__":
    main()

